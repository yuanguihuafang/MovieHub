from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
BASE = next(ROOT.glob("*2026"))
SRC = next(BASE.glob("*7章版.docx"))
OUT = SRC.with_name(f"{SRC.stem}_polished_final.docx")
REPORT = BASE / "thesis_7_final_report.md"


def set_run_font(run, size: float = 12, east: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)


def replace_para(p, text: str) -> None:
    p.clear()
    if not text:
        return
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if not text.strip().startswith(("第 ", "图 ", "表 ")):
        p.paragraph_format.first_line_indent = Pt(24)


def delete_para(p) -> None:
    el = p._element
    el.getparent().remove(el)
    p._p = p._element = None


def merge_adjacent_citations(text: str) -> str:
    pattern = re.compile(r"\[([0-9,\s]+)\]\[([0-9,\s]+)\]")
    old = None
    while old != text:
        old = text
        text = pattern.sub(lambda m: f"[{m.group(1).replace(' ', '')},{m.group(2).replace(' ', '')}]", text)
    return text


def main() -> None:
    if OUT.exists():
        OUT.unlink()
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    org_replacements = {
        "第1章 绪论": "第 1 章 绪论：阐述研究背景与意义，综述推荐系统、知识图谱表示学习、多模态融合与检索增强生成的研究现状，说明本文研究内容与贡献。",
        "第2章 关键技术介绍": "第 2 章 系统需求分析：从功能性需求和非功能性需求两方面建立需求模型，明确普通用户、管理员和推荐服务的需求边界。",
        "第 2 章 系统需求分析": "第 3 章 系统总体设计：给出系统总体架构、功能模块划分、数据库结构、接口设计与推荐流水线方案。",
        "第 3 章 系统总体设计": "第 4 章 系统详细设计：详细描述核心业务模块、数据库表结构、数据字典、推荐算法流程与缓存策略。",
        "第 4 章 系统详细设计": "第 5 章 系统功能实现：说明后端 API、推荐引擎、前端页面、数据库迁移、缓存和部署相关实现。",
        "第 5 章 系统功能实现": "第 6 章 系统测试：开展功能测试、推荐流程测试、知识图谱链路预测实验和推荐离线评估。",
        "第 6 章 系统测试": "第 7 章 总结与展望：总结系统建设与实验成果，分析不足并提出后续改进方向。",
    }

    replacements = {
        "KEY WORDS: Movie recommendation,Knowledge graph,Link prediction,Mixture of Experts,RAG,FastAPI": "KEY WORDS: Movie recommendation, Knowledge graph, Link prediction, Mixture of Experts, RAG, FastAPI",
        "用户可维护个人资料（昵称、偏好类型 preferred_genres 等）与修改密码。": "用户可维护偏好类型 preferred_genres 并修改密码。",
        "未配置 Redis 时自动降级为进程内内存缓存。": "未配置 Redis 时自动降级为进程内缓存。",
        "推荐引擎通过三层 Redis 缓存（KG 推理、RAG 检索、全量推荐）降低重复推理耗时。": "推荐引擎在配置 Redis 时使用 KG 推理、RAG 检索、全量推荐三层缓存；未配置 Redis 时降级为进程内缓存和磁盘缓存，以降低重复推理耗时并保证基础可用性。",
        "系统支持Redis可选缓存、进程内缓存与磁盘缓存的多层降级机制": "系统支持 Redis 可选缓存、进程内缓存与磁盘缓存的多层降级机制",
        "FastAPI+MySQL+ChromaDB+PyTorch": "FastAPI + MySQL + ChromaDB + PyTorch",
        "Vue 3+TypeScript+Vite+ Element Plus": "Vue 3 + TypeScript + Vite + Element Plus",
        "ChromaDB向量空间": "ChromaDB 向量空间",
        "DB15K中文别名词典": "DB15K 中文别名词典",
        "RAG辅助": "RAG 辅助",
        "HitRate@K、MRR": "HitRate@K、MRR、Coverage",
    }

    changed = 0
    to_delete = []
    for p in doc.paragraphs:
        raw = p.text
        stripped = raw.strip()
        if stripped == "第 7 章 总结与展望":
            to_delete.append(p)
            changed += 1
            continue
        new = raw
        if stripped in org_replacements:
            new = org_replacements[stripped]
        for src, dst in replacements.items():
            new = new.replace(src, dst)
        new = merge_adjacent_citations(new)
        if new != raw:
            replace_para(p, new)
            changed += 1

    for p in to_delete:
        delete_para(p)

    doc.save(OUT)

    newdoc = Document(OUT)
    text = "\n".join(p.text for p in newdoc.paragraphs) + "\n" + "\n".join(
        " ".join(c.text for r in t.rows for c in r.cells) for t in newdoc.tables
    )
    bad_terms = ["景点", "旅游线路", "酒店查询", "附近美食", "足迹", "AI行程", "spot_name", "footprint", "昵称"]
    remaining = {term: text.count(term) for term in bad_terms if text.count(term)}
    adjacent = re.findall(r"\[[0-9,\s]+\]\[[0-9,\s]+\]", text)
    refs = sorted(set(int(x) for x in re.findall(r"\[(\d+)\]", text) if int(x) < 100))
    org_section = []
    capture = False
    for p in newdoc.paragraphs:
        t = p.text.strip()
        if t == "本文组织结构":
            capture = True
            continue
        if capture:
            if t.startswith("第 2 章 系统需求分析"):
                org_section.append(t)
            elif t.startswith("第 ") or t.startswith("第1章"):
                org_section.append(t)
            if len(org_section) >= 7:
                break

    REPORT.write_text(
        f"""# 7章最终版继续完善报告

生成时间：2026-05-24

- 基准文件：`{SRC}`
- 输出文件：`{OUT}`

## 本次修改

1. 以用户指定的 `_7章版.docx` 为唯一基准重新处理。
2. 修正“本文组织结构”小节：去掉多余的“第2章关键技术介绍”，使其与 7 章正文结构一致。
3. 合并连续参考文献标注，例如 `[6,7][11]` 改为 `[6,7,11]`。
4. 修正英文关键词的逗号空格。
5. 修正“昵称”等与真实系统不一致的表述。
6. 统一 Redis 为可选缓存层，未配置时降级为进程内缓存/磁盘缓存。
7. 对部分技术栈表达补充空格，增强论文排版可读性。

## 自动复查

- 修改段落数：{changed}
- 段落数：{len(newdoc.paragraphs)}
- 表格数：{len(newdoc.tables)}
- 图片数：{len(newdoc.inline_shapes)}
- 旧模板/不一致关键词残留：{remaining}
- 连续引用残留：{adjacent}
- 检测到的参考文献编号：{refs}
- 组织结构抽样：{org_section}

## 最后人工步骤

1. 在 Word 中打开输出文件并更新目录。
2. 检查页眉页脚、封面信息是否符合学院最终提交模板。
3. 若学校要求 Word 原生交叉引用域，请使用“引用 -> 交叉引用”替换图表文字引用。
""",
        encoding="utf-8",
    )
    print(OUT)
    print(REPORT)
    print("changed", changed, "paragraphs", len(newdoc.paragraphs), "tables", len(newdoc.tables), "images", len(newdoc.inline_shapes))
    print("remaining", remaining, "adjacent", adjacent)


if __name__ == "__main__":
    main()
