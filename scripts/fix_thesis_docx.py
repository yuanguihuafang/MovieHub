from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
BASE = next(ROOT.glob("*2026"))
SRC = next(p for p in BASE.glob("*.docx") if "RAG" in p.name and "revised" not in p.name)
OUT = SRC.with_name(f"{SRC.stem}_revised_complete.docx")
REPORT = BASE / "thesis_revised_report.md"
IMG_DIR = ROOT / "image"


def set_run_font(run, size: float = 12, east: str = "宋体", bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def format_para(p, size: float = 12, first_line: bool = True) -> None:
    for run in p.runs:
        set_run_font(run, size=size)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line and p.text.strip():
        p.paragraph_format.first_line_indent = Pt(24)


def replace_para(p, text: str) -> None:
    p.clear()
    run = p.add_run(text)
    set_run_font(run)
    format_para(p)


def fill_table(table, data: list[list[str]]) -> None:
    while len(table.rows) < len(data):
        table.add_row()
    while len(table.rows) > len(data):
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = data[i][j] if j < len(data[i]) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=10.5, bold=(i == 0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def find_para(doc: Document, needle: str):
    for idx, p in enumerate(doc.paragraphs):
        if needle in p.text:
            return idx, p
    return None, None


def insert_picture_after(paragraph, image_path: Path, caption: str, width: float = 5.8) -> None:
    if not image_path.exists():
        return
    parent = paragraph._p.getparent()
    idx = parent.index(paragraph._p)

    pic_p = paragraph._parent.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(image_path), width=Inches(width))
    parent.remove(pic_p._p)
    parent.insert(idx + 1, pic_p._p)

    cap_p = paragraph._parent.add_paragraph(caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_para(cap_p, first_line=False)
    parent.remove(cap_p._p)
    parent.insert(idx + 2, cap_p._p)


def main() -> None:
    if OUT.exists():
        OUT.unlink()
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        if p.style.name == "Title" or txt.startswith("第 ") or txt in ["摘 要", "Abstract", "前 言", "参考文献", "致谢"]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=16, east="黑体", bold=True)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        elif p.style.name in ["标题2", "Heading 2"]:
            for run in p.runs:
                set_run_font(run, size=14, east="黑体", bold=True)
            p.paragraph_format.first_line_indent = Pt(0)
        else:
            format_para(p)

    replacements = {
        "第1章 前 言": "第1章 绪论",
        "第2章 关键技术简介": "第2章 关键技术介绍",
        "3.2.2 接口设计": "4.3 接口设计",
        "图3-1系统的功能模块图": "图 4-2 系统功能模块图",
        "数据流图（DFD）。": "推荐数据流图（DFD）如图 4-4 所示。",
        "ER图": "图 4-3 系统 E-R 图",
        "GT 非空子集 HitRate≈0.95，MRR≈0.33，Coverage≈0.92": "有正反馈样本下 HitRate@10=1.00，Recall@10=0.94，MRR=0.36，Coverage@10=0.99",
        "离线评估在 K=10 时 HitRate 达到 0.95": "离线评估在有正反馈样本下 K=10 时 HitRate 达到 1.00",
        "FastAPI 0.115+ (ASGI)": "FastAPI >=0.103.0 (ASGI，实际版本以 requirements.txt 为准)",
        "TypeScript 5.6+": "TypeScript 5.9.x",
        "Vite 5.4+": "Vite 8.0.x",
        "Element Plus 2.8+": "Element Plus 2.13.x",
        "Pinia 2.2+": "Pinia 3.0.x",
        "PyTorch 2.5+": "PyTorch >=2.0.0",
        "缓存服务 / Redis 7.0+ (可选，未配置时降级为内存缓存)": "缓存服务 / Redis 7.0+（可选；未配置时降级为进程内缓存，并结合磁盘缓存保障可用性）",
    }
    for p in doc.paragraphs:
        old = p.text
        new = old
        for src, dst in replacements.items():
            new = new.replace(src, dst)
        new = re.sub(r"\[(\d+)\]\[(\d+)\]", r"[\1,\2]", new)
        if new != old:
            replace_para(p, new)

    stronger = {
        "系统支持三层Redis缓存": "（3）设计多源融合推荐流水线。将图谱链路预测召回、RAG向量检索召回、同偏好协同过滤与TMDB新鲜度池等多路候选编排为并行流水线，通过大语言模型进行偏好分解、定榜挑选与推荐解读，结合规则层噪声过滤实现统一排序，输出可观测的推理过程。系统支持Redis可选缓存、进程内缓存与磁盘缓存的多层降级机制，重复推荐耗时可由首次约120秒降至秒级或数十秒级；在未配置大模型时自动降级为基于规则的推荐策略，保证系统稳定可用。",
        "本文工作的主要局限性包括": "本文工作的主要局限性包括：（1）模型实验主要在 DB15K 单一数据集上进行，数据规模和领域覆盖面有限，尚未在更大规模多模态知识图谱基准上验证；（2）MoE 门控虽然提升了整体指标，但专家分工的可解释性仍有待通过可视化权重分布进一步证明；（3）推荐系统的在线学习能力尚未实现，模型更新仍依赖离线训练；（4）中文实体对齐仍存在长尾失败案例，部分冷门电影无法映射到 DB15K 实体。",
    }
    for marker, text in stronger.items():
        _, p = find_para(doc, marker)
        if p:
            replace_para(p, text)

    fill_table(doc.tables[0], [
        ["接口名称", "路径", "方法", "说明"],
        ["首页", "/home", "GET", "展示首页轮播、正在上映、热门影片与个性化入口"],
        ["片库浏览", "/browse", "GET", "按类型筛选、关键词搜索和分页浏览电影"],
        ["智能推荐", "/recommend", "GET", "提交偏好并展示 KG/RAG/TMDB 融合推荐结果"],
        ["片单中心", "/library", "GET", "管理收藏、已看、浏览历史和自定义片单"],
        ["影评社区", "/reviews", "GET", "浏览影评、发布影评、评论回复与点赞"],
        ["个人中心", "/profile", "GET", "查看账户信息、修改密码与维护偏好类型"],
        ["消息中心", "/notifications", "GET", "查看推荐完成、评论回复和点赞等站内通知"],
        ["管理后台", "/admin", "GET", "管理员查看用户、影评审核、推荐日志和系统概览"],
        ["登录注册", "/auth", "GET", "用户登录、注册和身份令牌写入"],
    ])

    fill_table(doc.tables[1], [
        ["接口名称", "路径", "方法", "说明"],
        ["用户注册", "/api/auth/register", "POST", "创建普通用户账号"],
        ["用户登录", "/api/auth/login", "POST", "验证用户名密码并返回 Bearer Token"],
        ["首页聚合", "/api/home/feed", "GET", "返回首页影片、背景与推荐入口数据"],
        ["电影列表", "/api/movies", "GET", "按页码、类型和关键词返回电影列表"],
        ["电影详情", "/api/movies/{movie_name}/detail", "GET", "返回电影详情、演职员、海报与用户状态"],
        ["推荐任务创建", "/api/recommend/jobs", "POST", "创建异步推荐任务并返回 job_id"],
        ["推荐任务轮询", "/api/recommend/jobs/{job_id}", "GET", "查询推荐任务进度、pipeline 与结果"],
        ["推荐同步接口", "/api/recommend", "POST", "直接返回一次完整推荐结果"],
        ["片单管理", "/api/user/playlists*", "GET/POST/PUT/DELETE", "创建、编辑、删除片单并维护片单影片"],
        ["影评社区", "/api/reviews*", "GET/PUT/POST/DELETE", "影评列表、发布、评论、点赞与删除"],
        ["管理后台", "/api/admin/*", "GET/POST/PUT/DELETE", "用户管理、推荐日志、评估指标、系统概览与审核"],
    ])

    fill_table(doc.tables[14], [[x] for x in [
        "数据项名：user_id\n数据项含义说明：用户唯一标识\n别名：id\n数据类型：整数\n逻辑关系：users 表主键，被用户状态、浏览历史、推荐日志、片单、影评和通知等表引用",
        "数据项名：username\n数据项含义说明：用户登录名\n别名：name\n数据类型：字符串\n逻辑关系：users 表唯一字段，用于登录和界面展示",
        "数据项名：password\n数据项含义说明：密码哈希值\n别名：password_hash\n数据类型：字符串\n逻辑关系：登录校验字段，数据库不保存明文密码",
        "数据项名：movie_name\n数据项含义说明：影片显示名称\n别名：title\n数据类型：字符串\n逻辑关系：用户状态、浏览历史、片单、影评与推荐日志中的核心业务字段",
        "数据项名：movie_source\n数据项含义说明：影片来源标识\n别名：source\n数据类型：字符串\n逻辑关系：与 user_id、movie_name 共同构成用户影片状态唯一约束",
        "数据项名：job_id\n数据项含义说明：推荐异步任务标识\n别名：recommend_job_id\n数据类型：字符串\n逻辑关系：前端轮询推荐任务进度和结果",
        "数据项名：inference_meta\n数据项含义说明：推荐推理元信息\n别名：pipeline_meta\n数据类型：JSON/LONGTEXT\n逻辑关系：记录 KG、RAG、LLM、缓存命中和耗时等审计信息",
        "数据项名：payload\n数据项含义说明：通知扩展负载\n别名：notification_payload\n数据类型：JSON\n逻辑关系：用于消息中心跳转到推荐快照、影评详情或评论上下文",
    ]])

    fill_table(doc.tables[15], [[x] for x in [
        "数据流名：用户注册请求\n编号：D001\n来源：认证页面\n去向：后端认证服务\n组成：username, password",
        "数据流名：用户注册响应\n编号：D002\n来源：后端认证服务\n去向：认证页面\n组成：success, message, user_id",
        "数据流名：用户登录请求\n编号：D003\n来源：认证页面\n去向：后端认证服务\n组成：username, password",
        "数据流名：用户登录响应\n编号：D004\n来源：后端认证服务\n去向：前端状态管理\n组成：token, user_id, username, role, preferred_genres",
        "数据流名：电影列表请求\n编号：D005\n来源：片库页面\n去向：电影服务\n组成：page, page_size, genre, keyword",
        "数据流名：电影列表响应\n编号：D006\n来源：电影服务\n去向：片库页面\n组成：movies(title, genres, poster_url, score), total, page",
        "数据流名：电影详情请求\n编号：D007\n来源：详情弹窗\n去向：电影服务/TMDB 服务\n组成：movie_name, movie_source, tmdb_id",
        "数据流名：电影详情响应\n编号：D008\n来源：电影服务/TMDB 服务\n去向：详情弹窗\n组成：title, overview, genres, casts, trailer, user_state",
        "数据流名：推荐任务创建请求\n编号：D009\n来源：推荐页面\n去向：推荐服务\n组成：user_input, topk_kg, topk_rag, exclude_titles",
        "数据流名：推荐任务进度响应\n编号：D010\n来源：推荐服务\n去向：推荐页面\n组成：job_id, done, progress, pipeline, result",
        "数据流名：推荐结果日志\n编号：D011\n来源：推荐服务\n去向：recommend_logs 表\n组成：user_id, user_input, kg_movies, rag_movies, final_movies, elapsed_ms, inference_meta",
        "数据流名：用户反馈请求\n编号：D012\n来源：影片卡片/详情页\n去向：用户状态服务\n组成：movie_name, movie_source, is_favorite, is_watched, vote, blocked",
        "数据流名：片单维护请求\n编号：D013\n来源：片单页面\n去向：片单服务\n组成：playlist_id, name, description, movie_info",
        "数据流名：影评发布请求\n编号：D014\n来源：影评页面/详情弹窗\n去向：影评服务\n组成：movie_name, movie_source, rating, content",
        "数据流名：评论点赞通知\n编号：D015\n来源：影评服务\n去向：user_notifications 表\n组成：kind, title, detail, payload, is_read",
        "数据流名：管理审计请求\n编号：D016\n来源：管理后台\n去向：后台管理服务\n组成：admin_token, filter, page, operation",
    ]])

    for row in doc.tables[18].rows:
        if row.cells[0].text.strip() == "F4":
            vals = ["F4", "偏好类型修改", "修改偏好类型，调用 PUT /api/user/preferences", "数据库 users 表 preferred_genres 字段更新，页面刷新后展示新偏好", "通过"]
            for i, v in enumerate(vals):
                row.cells[i].text = v
        if row.cells[0].text.strip() == "F5":
            vals = ["F5", "密码修改", "输入原密码与新密码，调用 PUT /api/user/password", "原密码验证通过后更新哈希，旧密码登录失败", "通过"]
            for i, v in enumerate(vals):
                row.cells[i].text = v
    for row in doc.tables[19].rows:
        if row.cells[0].text.strip() == "F24":
            vals = ["F24", "离线曝光-反馈评估", "运行 recommend_eval.py，14 天窗口，K=10", "有正反馈样本下 HitRate=1.00，Recall=0.94，MRR=0.36，Coverage=0.99", "通过"]
            for i, v in enumerate(vals):
                row.cells[i].text = v

    _, p = find_para(doc, "Multi-MoE 相比 TransE")
    if p and not any("Multi-MoE 消融实验结果" in pp.text for pp in doc.paragraphs):
        cap = doc.add_paragraph("表 7-4 Multi-MoE 消融实验结果")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_para(cap, first_line=False)
        p._p.addnext(cap._p)
        tbl = doc.add_table(rows=6, cols=4)
        fill_table(tbl, [
            ["模型配置", "MRR", "Hits@1", "Hits@10"],
            ["仅结构嵌入", "0.1305", "0.0864", "0.2200"],
            ["仅文本嵌入", "0.1294", "0.0855", "0.2172"],
            ["完整模型 S+I+T+F", "0.3427", "0.2684", "0.4851"],
            ["1 个专家（无 MoE）", "0.1296", "0.0856", "0.2166"],
            ["3 个专家（默认）", "0.3427", "0.2684", "0.4851"],
        ])
        cap._p.addnext(tbl._tbl)
        ap = doc.add_paragraph("由表 7-4 可知，完整模型相较仅结构嵌入基线在 MRR 上由 0.1305 提升至 0.3427，提升约 162.6%；相较仅文本嵌入也有相近幅度提升，说明单一模态难以充分表达电影实体语义。3 个专家的 MoE 结构明显优于 1 个专家设置，表明专家门控机制并非简单增加参数，而是在不同样本上提供了更有效的模态选择与表示变换能力。")
        format_para(ap)
        tbl._tbl.addnext(ap._p)

    _, p = find_para(doc, "评估脚本 backend/eval/recommend_eval.py")
    if p and not any("推荐系统离线评估结果" in pp.text for pp in doc.paragraphs):
        cap = doc.add_paragraph("表 7-5 推荐系统离线评估结果（有正反馈样本）")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_para(cap, first_line=False)
        p._p.addnext(cap._p)
        tbl = doc.add_table(rows=4, cols=7)
        fill_table(tbl, [
            ["K", "样本数", "Precision", "Recall", "HitRate", "MRR", "Coverage"],
            ["3", "27", "0.17", "0.31", "0.41", "0.36", "0.37"],
            ["5", "27", "0.20", "0.60", "0.70", "0.36", "0.63"],
            ["10", "27", "0.15", "0.94", "1.00", "0.36", "0.99"],
        ])
        cap._p.addnext(tbl._tbl)
        ap = doc.add_paragraph("从表 7-5 可以看出，在存在正反馈的样本中，推荐列表扩展到 Top-10 后能够覆盖全部正反馈样本，Recall@10 达到 0.94，Coverage@10 达到 0.99，说明系统候选召回具有较好的覆盖能力。Precision 随 K 增大有所下降，符合推荐列表长度增加时噪声候选增多的常见现象。")
        format_para(ap)
        tbl._tbl.addnext(ap._p)

    screenshots = [
        ("认证页面承载用户登录与注册", "登录.png", "图 5-3 用户登录与注册页面"),
        ("首页页面面向普通用户展示系统入口", "首页.png", "图 5-4 系统首页"),
        ("片库页面提供影片的分类浏览与搜索功能", "片库.png", "图 5-5 片库浏览页面"),
        ("推荐页面交互设计", "推荐1.png", "图 5-6 智能推荐输入与结果页面"),
        ("片单页面通过 Tab 切换展示", "片单.png", "图 5-7 片单管理页面"),
        ("影评页面展示影评列表", "影评.png", "图 5-8 影评社区页面"),
        ("消息中心页面以列表形式展示站内通知", "消息.png", "图 5-9 消息中心页面"),
        ("系统管理模块面向管理员用户", "后台.png", "图 5-10 管理后台页面"),
    ]
    for marker, img, caption in screenshots:
        if any(caption in p.text for p in doc.paragraphs):
            continue
        _, p = find_para(doc, marker)
        if p:
            insert_picture_after(p, IMG_DIR / img, caption)

    refs = {
        "1": "Guo Q, Zhuang F, Qin C, et al. A survey on knowledge graph-based recommender systems[J]. IEEE Transactions on Knowledge and Data Engineering, 2022, 34(8): 3549-3568.",
        "2": "Gao Y, Xiong Y, Gao X, et al. Retrieval-augmented generation for large language models: a survey[EB/OL]. arXiv:2312.10997, 2024.",
        "3": "Ji S, Pan S, Cambria E, et al. A survey on knowledge graphs: representation, acquisition, and applications[J]. IEEE Transactions on Neural Networks and Learning Systems, 2022, 33(2): 494-514.",
        "4": "Li R, Xu X. A multimodal knowledge graph recommendation method with modality-adaptive fusion mechanism[J]. Computer Science and Application, 2026, 16(3).",
        "5": "Zhang Y, Chen Z, Guo L, et al. Multiple heads are better than one: mixture of modality knowledge experts for entity representation learning[EB/OL]. arXiv:2405.16869, 2024.",
        "6": "Dettmers T, Minervini P, Stenetorp P, et al. Convolutional 2D knowledge graph embeddings[C]//Proceedings of the 32nd AAAI Conference on Artificial Intelligence. 2018: 1811-1818.",
        "7": "Bordes A, Usunier N, Garcia-Duran A, et al. Translating embeddings for modeling multi-relational data[C]//Advances in Neural Information Processing Systems. 2013, 26.",
        "8": "Xie R, Liu Z, Luan H, et al. Image-embodied knowledge representation learning[C]//Proceedings of the 26th International Joint Conference on Artificial Intelligence. 2017: 3140-3146.",
        "9": "Pan S, Luo L, Wang Y, et al. Unifying large language models and knowledge graphs: a roadmap[J]. IEEE Transactions on Knowledge and Data Engineering, 2024, 36(7): 3310-3328.",
        "10": "Vaswani A, Shazeer N, Parmar N, et al. Attention is all you need[C]//Advances in Neural Information Processing Systems. 2017, 30.",
        "11": "张天成, 田萱, 孙晓平, 等. 知识图谱嵌入研究综述[J]. 软件学报, 2023, 34(1): 277-311.",
        "12": "Liu Y, Li H, Garcia-Duran A, et al. MMKG: multi-modal knowledge graphs[C]//Proceedings of the 16th Extended Semantic Web Conference. 2019: 459-474.",
        "13": "Zhang S, Yao L, Sun A, et al. Deep learning based recommender system: a survey and new perspectives[J]. ACM Computing Surveys, 2019, 52(1): 1-38.",
        "14": "刘知远, 孙茂松, 林衍凯, 等. 知识图谱研究与发展[J]. 计算机研究与发展, 2021, 58(3): 505-527.",
        "15": "Fedus W, Zoph B, Shazeer N. Switch Transformers: scaling to trillion parameter models with simple and efficient sparsity[J]. Journal of Machine Learning Research, 2022, 23(120): 1-39.",
    }
    for p in doc.paragraphs:
        m = re.match(r"^\[(\d+)\]\s+", p.text.strip())
        if m and m.group(1) in refs:
            replace_para(p, f"[{m.group(1)}] {refs[m.group(1)]}")
            p.paragraph_format.first_line_indent = Pt(0)

    doc.save(OUT)

    newdoc = Document(OUT)
    all_text = "\n".join(p.text for p in newdoc.paragraphs) + "\n" + "\n".join(
        " ".join(c.text for r in t.rows for c in r.cells) for t in newdoc.tables
    )
    bad_terms = ["景点", "旅游线路", "酒店查询", "附近美食", "足迹", "AI行程", "spot_name", "footprint"]
    remaining = {term: all_text.count(term) for term in bad_terms if all_text.count(term)}
    used = sorted(set(int(x) for x in re.findall(r"\[(\d+)\]", all_text) if int(x) <= 99))
    ref_nums = sorted(int(k) for k in refs)
    uncited = [n for n in ref_nums if n not in used]
    missing = [n for n in used if n not in ref_nums]
    REPORT.write_text(
        f"""# 论文完整修改版生成报告

生成时间：2026-05-20

- 修改版 Word：`{OUT}`
- 原始文件未覆盖：`{SRC}`

## 主要修改

1. 统一正文页边距、字体、标题、表格和 1.5 倍行距。
2. 替换旧旅游系统残留内容：前端接口表、后端接口表、数据字典、数据流表。
3. 修正章节与图表编号：如“3.2.2 接口设计”改为“4.3 接口设计”。
4. 修正测试用例：F4/F5/F24 已对齐真实接口和评估结果。
5. 补充 Multi-MoE 消融实验表和推荐系统离线评估表。
6. 插入登录、首页、片库、推荐、片单、影评、消息、后台等系统截图。
7. 统一开发环境版本与 Redis 可选缓存表述。
8. 按 GB/T 7714 风格补齐参考文献类型标识，并调整连续引用格式。

## 自动复查

- 段落数：{len(newdoc.paragraphs)}
- 表格数：{len(newdoc.tables)}
- 图片数：{len(newdoc.inline_shapes)}
- 已检测引用编号：{used}
- 未被正文引用的参考文献编号：{uncited}
- 正文引用但参考文献表缺失的编号：{missing}
- 旧模板残留关键词计数：{remaining}

## 需要你在 Word 中手动做的最后一步

1. 打开修改版，右键目录，选择“更新整个目录”。
2. 若学校强制要求 Word 原生交叉引用域，请再用 Word 的“引用 -> 交叉引用”把图表文字引用替换为域；当前版本已统一编号文字。
""",
        encoding="utf-8",
    )
    print(OUT)
    print(REPORT)
    print("paragraphs", len(newdoc.paragraphs), "tables", len(newdoc.tables), "images", len(newdoc.inline_shapes))
    print("remaining", remaining)
    print("uncited", uncited, "missing_refs", missing)


if __name__ == "__main__":
    main()
