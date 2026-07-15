from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
BASE = next(ROOT.glob("*2026"))
SRC = next(BASE.glob("*RAG*revised_complete.docx"))
OUT = SRC.with_name(f"{SRC.stem}_polished_v2.docx")
REPORT = BASE / "thesis_polished_v2_report.md"


def set_run_font(run, size: float = 12, east: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)


def replace_para(p, text: str) -> None:
    p.clear()
    run = p.add_run(text)
    set_run_font(run)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if text.strip() and not text.strip().startswith(("第 ", "图 ", "表 ")):
        p.paragraph_format.first_line_indent = Pt(24)


def merge_adjacent_citations(text: str) -> str:
    # [6,7][11] -> [6,7,11], applied repeatedly for longer chains.
    pattern = re.compile(r"\[([0-9,\s]+)\]\[([0-9,\s]+)\]")
    old = None
    while old != text:
        old = text
        text = pattern.sub(lambda m: f"[{m.group(1).replace(' ', '')},{m.group(2).replace(' ', '')}]", text)
    return text


def main() -> None:
    if OUT.exists():
        OUT.unlink()
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    replacements = {
        "本文共分为 7 章，组织结构如下：": "本文共分为 8 章，组织结构如下：",
        "第 2 章 关键技术简介：介绍系统涉及的核心技术，包括知识图谱嵌入与链接预测、多模态融合与 MoE 机制、检索增强生成与向量检索，以及 FastAPI、Vue 3 等工程技术栈。": "第 2 章 关键技术介绍：介绍系统涉及的核心技术，包括知识图谱嵌入与链接预测、多模态融合与 MoE 机制、检索增强生成与向量检索，以及 FastAPI、Vue 3 等工程技术栈。",
        "系统支持三层 Redis 缓存（KG 推理、RAG 检索、全量推荐），重复推荐耗时从约 120 秒降至约 25 秒。": "系统支持 Redis 可选缓存、进程内缓存与磁盘缓存的多层降级机制；当 Redis 可用时，可缓存 KG 推理、RAG 检索与全量推荐结果，重复推荐耗时可由首次约 120 秒降至秒级或数十秒级。",
        "MovieHub采用三层Redis缓存策略加速推荐：第一层缓存Multi-MoE知识图谱推理结果（以种子电影、关系类型与偏好类型为key），第二层缓存RAG向量检索与LLM选片结果（以查询文本与类型为key），第三层缓存完整推荐结果（以用户ID": "MovieHub 将 Redis 作为可选缓存层使用；配置 MOVIEHUB_REDIS_URL 时，系统采用三层缓存策略加速推荐：第一层缓存 Multi-MoE 知识图谱推理结果（以种子电影、关系类型与偏好类型为 key），第二层缓存 RAG 向量检索与 LLM 选片结果（以查询文本与类型为 key），第三层缓存完整推荐结果（以用户 ID",
        "未配置Redis时自动降级为内存缓存": "未配置 Redis 时自动降级为进程内缓存",
        "Redis多层缓存": "Redis 可选多层缓存",
        "三层 Redis 缓存策略": "Redis 可选三层缓存策略",
        "三层Redis缓存策略": "Redis 可选三层缓存策略",
    }

    changed = 0
    for p in doc.paragraphs:
        old = p.text
        new = old
        for src, dst in replacements.items():
            new = new.replace(src, dst)
        new = merge_adjacent_citations(new)
        if new != old:
            replace_para(p, new)
            changed += 1

    # Fix English keyword spacing.
    for p in doc.paragraphs:
        old = p.text
        if old.startswith("KEY WORDS:"):
            new = "KEY WORDS: Movie recommendation, Knowledge graph, Link prediction, Mixture of Experts, RAG, FastAPI"
            replace_para(p, new)
            changed += 1

    doc.save(OUT)

    newdoc = Document(OUT)
    text = "\n".join(p.text for p in newdoc.paragraphs) + "\n" + "\n".join(
        " ".join(c.text for r in t.rows for c in r.cells) for t in newdoc.tables
    )
    bad_terms = ["景点", "旅游线路", "酒店查询", "附近美食", "足迹", "AI行程", "spot_name", "footprint"]
    remaining = {term: text.count(term) for term in bad_terms if text.count(term)}
    adjacent = re.findall(r"\[[0-9,\s]+\]\[[0-9,\s]+\]", text)
    refs = sorted(set(int(x) for x in re.findall(r"\[(\d+)\]", text) if int(x) < 100))
    has_wrong_chapter_count = "本文共分为 7 章" in text

    REPORT.write_text(
        f"""# 论文继续完善报告 v2

生成时间：2026-05-24

- 基准文件：`{SRC}`
- 输出文件：`{OUT}`

## 本次继续完善

1. 修正“本文共分为 7 章”与实际 8 章结构不一致的问题。
2. 将“第 2 章 关键技术简介”统一为“第 2 章 关键技术介绍”。
3. 合并连续参考文献标注，例如 `[6,7][11]` 改为 `[6,7,11]`。
4. 统一 Redis 表述为“可选缓存层，未配置时降级为进程内缓存/磁盘缓存”。
5. 修正英文关键词逗号后的空格格式。

## 自动复查

- 修改段落数：{changed}
- 段落数：{len(newdoc.paragraphs)}
- 表格数：{len(newdoc.tables)}
- 图片数：{len(newdoc.inline_shapes)}
- 旧旅游模板关键词残留：{remaining}
- 连续引用残留：{adjacent}
- 检测到的参考文献编号：{refs}
- 是否仍存在“本文共分为 7 章”：{has_wrong_chapter_count}

## 仍需人工确认

1. 用 Word 打开输出文件后，更新目录。
2. 若学校要求必须使用 Word 原生交叉引用域，请用 Word 的交叉引用功能替换图表文字引用。
""",
        encoding="utf-8",
    )
    print(OUT)
    print(REPORT)
    print("changed", changed, "paragraphs", len(newdoc.paragraphs), "tables", len(newdoc.tables), "images", len(newdoc.inline_shapes))
    print("remaining", remaining, "adjacent", adjacent, "wrong_chapter_count", has_wrong_chapter_count)


if __name__ == "__main__":
    main()
